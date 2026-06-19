"""
PDF and Word Document Generator Module
Generates professional quotation documents in PDF and Word formats
"""

import os
import re
from datetime import datetime
from xhtml2pdf import pisa
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any


def _safe_name(value: str) -> str:
    """Sanitize a string for use in a filename (remove/replace unsafe chars)."""
    value = value.strip()
    # Replace spaces and special characters with underscores
    value = re.sub(r'[^\w\-]', '_', value)
    # Collapse multiple underscores
    value = re.sub(r'_+', '_', value)
    return value or 'unknown'


class DocumentGenerator:
    """Generate quotation documents in PDF and Word formats"""
    
    def __init__(self, output_folder: str = 'outputs'):
        self.output_folder = output_folder
        os.makedirs(output_folder, exist_ok=True)
    
    def generate_pdf(self, quotation_data: Dict[str, Any], template_type: str,
                     company_name: str = '', full_name: str = '') -> str:
        """
        Generate PDF from quotation data
        
        Args:
            quotation_data: Dictionary containing quotation information
            template_type: 'type1' or 'type2'
            company_name: Company name for the filename
            full_name: Full name of the user for the filename
            
        Returns:
            Path to generated PDF file
        """
        html_content = self._create_html_content(quotation_data, template_type)
        
        # Generate filename: {company_name}_{full_name}_{timestamp}.pdf
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        company_part = _safe_name(company_name) if company_name else 'company'
        name_part = _safe_name(full_name) if full_name else 'client'
        filename = f"{company_part}_{name_part}_{timestamp}.pdf"
        output_path = os.path.join(self.output_folder, filename)
        
        # Convert HTML to PDF using xhtml2pdf
        with open(output_path, 'wb') as pdf_file:
            pisa_status = pisa.CreatePDF(
                html_content,
                dest=pdf_file
            )
        
        if pisa_status.err:
            raise Exception(f"PDF generation failed with error code: {pisa_status.err}")
        
        return output_path
    
    def generate_word(self, quotation_data: Dict[str, Any], template_type: str,
                      company_name: str = '', full_name: str = '') -> str:
        """
        Generate Word document from quotation data
        
        Args:
            quotation_data: Dictionary containing quotation information
            template_type: 'type1' or 'type2'
            company_name: Company name for the filename
            full_name: Full name of the user for the filename
            
        Returns:
            Path to generated Word file
        """
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        if template_type == 'type1':
            self._create_word_type1(doc, quotation_data)
        else:
            self._create_word_type2(doc, quotation_data)
        
        # Generate filename: {company_name}_{full_name}_{timestamp}.docx
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        company_part = _safe_name(company_name) if company_name else 'company'
        name_part = _safe_name(full_name) if full_name else 'client'
        filename = f"{company_part}_{name_part}_{timestamp}.docx"
        output_path = os.path.join(self.output_folder, filename)
        
        doc.save(output_path)
        
        return output_path
    
    def _create_html_content(self, data: Dict[str, Any], template_type: str) -> str:
        """Create HTML content for PDF generation"""
        
        if template_type == 'type1':
            return self._create_html_type1(data)
        else:
            return self._create_html_type2(data)
    
    def _create_html_type1(self, data: Dict[str, Any]) -> str:
        """Create HTML for Type 1 template using Jinja2 template"""
        from jinja2 import Template
        
        # Load template file
        template_path = os.path.join(os.path.dirname(__file__), '..', 'templates', 'type1', 'type1_template.html')
        with open(template_path, 'r', encoding='utf-8') as f:
            template_content = f.read()
        
        # Create Jinja2 template
        template = Template(template_content)
        
        # Prepare data with defaults
        template_data = {
            'pan_number': data.get('pan_number', 'AGVPJ1503G'),
            'gstin': data.get('gstin', '27AGVPJ1503G1ZF'),
            'proposal_validity': data.get('proposal_validity', '10th Nov 2025'),
            'date': data.get('date', datetime.now().strftime('%d/%m/%Y')),
            'project_title': data.get('project_title', 'Proposal'),
            'scope_of_work': data.get('scope_of_work', []),
            'pricing_table': data.get('pricing_table', []),
            'terms_and_conditions': data.get('terms_and_conditions', []),
            'bank_name': data.get('bank_name', 'HDFC Bank'),
            'account_name': data.get('account_name', 'MJ Services'),
            'account_number': data.get('account_number', '50200012345678'),
            'ifsc_code': data.get('ifsc_code', 'HDFC0001234'),
            'micr_code': data.get('micr_code', '400240012'),
            'authorized_signatory_name': data.get('authorized_signatory_name', 'Ms. Mili Juneja'),
            'company_name': data.get('company_name', 'MJ Services'),
            'company_role': data.get('company_role', 'Authorized Zoho Channel Partner'),
            'logo_path': os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'static', 'logo.png')),
        }
        
        # Render template
        html = template.render(**template_data)
        
        return html

    
    def _create_html_type2(self, data: Dict[str, Any]) -> str:
        """Create HTML for Type 2 template using Jinja2 template"""
        from jinja2 import Template
        
        # Load template file
        template_path = os.path.join(os.path.dirname(__file__), '..', 'templates', 'type2', 'type2_template.html')
        with open(template_path, 'r', encoding='utf-8') as f:
            template_content = f.read()
        
        # Create Jinja2 template
        template = Template(template_content)
        
        # Prepare data with defaults
        template_data = {
            'project_title': data.get('project_title', 'Project Proposal'),
            'client_name': data.get('client_name', 'Valued Client'),
            'date': data.get('date', datetime.now().strftime('%d %B %Y')),
            'reference_number': data.get('reference_number', 'QT-2024-001'),
            'executive_summary': data.get('executive_summary', ''),
            'scope_of_work': data.get('scope_of_work', []),
            'pricing_table': data.get('pricing_table', []),
            'grand_total': data.get('grand_total', '0'),
            'additional_notes': data.get('additional_notes', 'Our proven methodology ensures successful project delivery with regular milestones and quality checkpoints.'),
            'logo_path': os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'static', 'logo.png')),
        }
        
        # Render template
        html = template.render(**template_data)
        
        return html

    
    def _create_word_type1(self, doc: Document, data: Dict[str, Any]):
        """Create Word document for Type 1 template"""
        
        # Logo placeholder
        p = doc.add_paragraph()
        logo_path = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'static', 'logo.png'))
        if os.path.exists(logo_path):
            p.add_run().add_picture(logo_path, width=Inches(1.5))
        else:
            p.add_run('[LOGO]').bold = True
        
        # Date
        doc.add_paragraph(f"Date: {data.get('date', datetime.now().strftime('%d/%m/%Y'))}")
        
        # Title
        title = doc.add_heading(data.get('project_title', 'Proposal'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Scope of Work
        doc.add_heading('Scope of Work/Deliverables', 1)
        for idx, item in enumerate(data.get('scope_of_work', []), 1):
            if '||' in item:
                parts = item.split('||')
                title_part = parts[0].strip()
                bullets = [b.strip() for b in parts[1:] if b.strip()]
                p = doc.add_paragraph()
                p.add_run(f"{idx}. {title_part}").bold = True
                for bullet in bullets:
                    bp = doc.add_paragraph(style='List Bullet')
                    bp.add_run(bullet)
                    bp.paragraph_format.left_indent = Inches(0.5)
            elif ':' in item:
                title_part = item.split(':')[0].strip()
                desc_part = item.split(':', 1)[1].strip()
                p = doc.add_paragraph()
                p.add_run(f"{idx}. {title_part}").bold = True
                bp = doc.add_paragraph(style='List Bullet')
                bp.add_run(desc_part)
                bp.paragraph_format.left_indent = Inches(0.5)
            else:
                p = doc.add_paragraph()
                p.add_run(f"{idx}. {item}").bold = True
        
        # Commercials
        doc.add_heading('Commercials:', 1)
        pricing_table = data.get('pricing_table', [])
        if pricing_table:
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Light Grid Accent 1'
            
            # Header row with yellow background
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Service Description'
            hdr_cells[1].text = 'Fees (INR)'
            
            # Data rows
            for item in pricing_table:
                row_cells = table.add_row().cells
                row_cells[0].text = str(item.get('description', ''))
                row_cells[1].text = f"₹{item.get('total_price', '')}"
        
        # Terms & Notes
        doc.add_heading('Terms & Notes', 1)
        for term in data.get('terms_and_conditions', []):
            doc.add_paragraph(term, style='List Bullet')
        
        # Bank Details
        doc.add_heading('Bank Details for Payment', 1)
        bank_info = [
            f"Bank Name: {data.get('bank_name', 'HDFC Bank')}",
            f"Account Name: {data.get('account_name', 'MJ Services')}",
            f"Account Number: {data.get('account_number', '50200012345678')}",
            f"IFSC Code: {data.get('ifsc_code', 'HDFC0001234')}",
            f"MICR Code: {data.get('micr_code', '400240012')}"
        ]
        for info in bank_info:
            doc.add_paragraph(info, style='List Bullet')
        
        # Company Info (PAN/GSTIN) - At the end
        doc.add_paragraph(f"PAN: {data.get('pan_number', 'AGVPJ1503G')}")
        doc.add_paragraph(f"GSTIN: {data.get('gstin', '27AGVPJ1503G1ZF')}")
        
        # Proposal Validity - At the end
        p = doc.add_paragraph()
        p.add_run('Proposal Validity').bold = True
        doc.add_paragraph(f"This proposal is valid until {data.get('proposal_validity', '10th Nov 2025')}. Kindly share your confirmation to initiate the process.")
        
        # Authorized Signatory - Last section
        doc.add_heading('Authorized Signatory', 1)
        p = doc.add_paragraph()
        p.add_run(data.get('authorized_signatory_name', 'Ms. Mili Juneja')).bold = True
        doc.add_paragraph(data.get('company_name', 'MJ Services'))
        doc.add_paragraph(data.get('company_role', 'Authorized Zoho Channel Partner'))

    
    def _create_word_type2(self, doc: Document, data: Dict[str, Any]):
        """Create Word document for Type 2 template"""
        
        # Date
        doc.add_paragraph(f"Date: {data.get('date', datetime.now().strftime('%d/%m/%Y'))}")
        
        # Title (underlined and centered)
        title = doc.add_heading(data.get('project_title', 'Proposal'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.underline = True
        
        # Concerned Person (if available)
        client_name = data.get('client_name', '')
        if client_name and client_name != 'Valued Client':
            p = doc.add_paragraph()
            p.add_run('Concerned Person: ').bold = True
            p.add_run(client_name)
        
        # Objective (if available)
        if data.get('executive_summary'):
            p = doc.add_paragraph()
            p.add_run('Objective: ').bold = True
            p.add_run(data.get('executive_summary'))
        
        # Scope of Work
        doc.add_heading('Scope of work:', 1)
        for idx, item in enumerate(data.get('scope_of_work', []), 1):
            doc.add_paragraph(item, style='List Number')
        
        # Commercials
        doc.add_heading('Commercials:', 1)
        pricing_table = data.get('pricing_table', [])
        if pricing_table and len(pricing_table) > 0:
            commercials_text = ', '.join([
                f"{item.get('description', '')}{'₹' + item.get('total_price', '') if item.get('total_price') else ''}"
                for item in pricing_table
            ])
            doc.add_paragraph(commercials_text)
        else:
            grand_total = data.get('grand_total', 'XXXXXX')
            doc.add_paragraph(f"One-time Fee of ₹{grand_total} + GST, Payment to be made in advance.")
        
        # Note / Terms
        doc.add_heading('Note:', 1)
        for idx, term in enumerate(data.get('terms_and_conditions', []), 1):
            doc.add_paragraph(term, style='List Number')
        
        # Proposal Validity
        doc.add_paragraph("This proposal is valid for this month. Request you to share your consent on the same.")
        
        # Signatory
        p = doc.add_paragraph()
        p.add_run('Signatory:').bold = True
        p = doc.add_paragraph()
        p.add_run(data.get('authorized_signatory_name', 'Ms. Mili Juneja')).bold = True
        doc.add_paragraph(data.get('company_name', 'MJ Services'))
        doc.add_paragraph(data.get('company_role', 'Authorized Zoho Channel Partner'))
        
        # Bank Account Details
        heading = doc.add_heading('Bank Account Details:', 1)
        for run in heading.runs:
            run.underline = True
        
        bank_details = [
            f"Bank Name: {data.get('bank_name', 'City Union Bank')}",
            f"Account Name: {data.get('account_name', 'MJ Services')}",
            f"Account Number: {data.get('account_number', '510909010018772')}",
            f"IFSC Code: {data.get('ifsc_code', 'CIUB0000178')}",
            f"MICR Code: {data.get('micr_code', '440054002')}",
            f"PAN Number: {data.get('pan_number', 'AGVPJ1503G')}",
            f"GSTIN: {data.get('gstin', '27AGVPJ1503G1ZF')}"
        ]
        for detail in bank_details:
            doc.add_paragraph(detail)

