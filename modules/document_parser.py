"""
Document Parser Module
Handles parsing of Word (.docx), PDF, Excel, and image files to extract requirement text
Enhanced with OCR and NLP entity extraction (optional)
"""

import os
import re
from docx import Document
import PyPDF2
from typing import Dict, Any, List, Optional

# Optional imports - gracefully handle if not installed
try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False
    print("Warning: pandas not installed. Excel support disabled.")

try:
    from modules.ocr_handler import OCRHandler
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    print("Warning: OCR dependencies not installed. OCR features disabled.")

try:
    from modules.entity_extractor import EntityExtractor
    NLP_AVAILABLE = True
except ImportError:
    NLP_AVAILABLE = False
    print("Warning: NLP dependencies not installed. Entity extraction disabled.")


class DocumentParser:
    """Parse Word, PDF, Excel, and image documents to extract text content"""
    
    def __init__(self, enable_ocr: bool = True, enable_nlp: bool = True):
        """
        Initialize document parser
        
        Args:
            enable_ocr: Enable OCR for scanned documents (if available)
            enable_nlp: Enable NLP entity extraction (if available)
        """
        self.supported_formats = ['.docx', '.doc', '.pdf']
        self.image_formats = {'.png', '.jpg', '.jpeg', '.tiff'}
        self.excel_formats = {'.xlsx', '.xls'}
        
        # Add Excel and image formats only if dependencies are available
        if PANDAS_AVAILABLE:
            self.supported_formats.extend(['.xlsx', '.xls'])
        
        if OCR_AVAILABLE:
            self.supported_formats.extend(['.png', '.jpg', '.jpeg', '.tiff'])
        
        # Initialize OCR handler
        self.ocr_handler = None
        if enable_ocr and OCR_AVAILABLE:
            try:
                self.ocr_handler = OCRHandler()
            except Exception as e:
                print(f"Warning: OCR initialization failed: {e}")
        
        # Initialize entity extractor
        self.entity_extractor = None
        if enable_nlp and NLP_AVAILABLE:
            try:
                self.entity_extractor = EntityExtractor()
            except Exception as e:
                print(f"Warning: NLP initialization failed: {e}")
    
    def parse_document(self, file_path: str) -> Dict[str, Any]:
        """
        Parse document and extract text content
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext in ['.docx', '.doc']:
            return self._parse_word_document(file_path)
        elif file_ext == '.pdf':
            return self._parse_pdf_document(file_path)
        elif file_ext in self.excel_formats:
            return self._parse_excel_document(file_path)
        elif file_ext in self.image_formats:
            return self._parse_image_document(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    def _parse_word_document(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from Word document
        
        Args:
            file_path: Path to .docx file
            
        Returns:
            Dictionary with extracted text and metadata
        """
        try:
            doc = Document(file_path)
            
            # Extract all paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
            
            # Extract tables
            tables_text = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables_text.append(table_data)
            
            full_text = '\n\n'.join(paragraphs)
            
            # Detect sections
            sections = self._detect_sections(full_text)
            
            # Extract entities if NLP is enabled
            structured_data = None
            if self.entity_extractor:
                structured_data = self.entity_extractor.extract_structured_data(full_text)
            
            return {
                'success': True,
                'text': full_text,
                'paragraphs': paragraphs,
                'tables': tables_text,
                'sections': sections,
                'structured_data': structured_data,
                'file_type': 'word',
                'filename': os.path.basename(file_path)
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'file_type': 'word',
                'filename': os.path.basename(file_path)
            }
    
    def _parse_pdf_document(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from PDF document (with OCR support for scanned PDFs)
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Dictionary with extracted text and metadata
        """
        try:
            # First try regular text extraction
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract text from all pages
                pages_text = []
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text.strip():
                        pages_text.append(text.strip())
                
                full_text = '\n\n'.join(pages_text)
                
                # Check if we got substantial text
                if len(full_text.strip()) < 50 and self.ocr_handler:
                    # Likely a scanned PDF, use OCR
                    ocr_result = self.ocr_handler.extract_text_from_pdf(file_path)
                    
                    if ocr_result.get('success'):
                        full_text = ocr_result.get('text', '')
                        
                        # Detect sections
                        sections = self._detect_sections(full_text)
                        
                        # Extract entities
                        structured_data = None
                        if self.entity_extractor:
                            structured_data = self.entity_extractor.extract_structured_data(full_text)
                        
                        return {
                            'success': True,
                            'text': full_text,
                            'pages': [full_text],
                            'page_count': ocr_result.get('page_count', 0),
                            'sections': sections,
                            'structured_data': structured_data,
                            'file_type': 'pdf',
                            'method': 'ocr',
                            'ocr_confidence': ocr_result.get('confidence', 0),
                            'filename': os.path.basename(file_path)
                        }
                
                # Regular text extraction succeeded
                sections = self._detect_sections(full_text)
                
                structured_data = None
                if self.entity_extractor:
                    structured_data = self.entity_extractor.extract_structured_data(full_text)
                
                return {
                    'success': True,
                    'text': full_text,
                    'pages': pages_text,
                    'page_count': len(pdf_reader.pages),
                    'sections': sections,
                    'structured_data': structured_data,
                    'file_type': 'pdf',
                    'method': 'text_extraction',
                    'filename': os.path.basename(file_path)
                }
                
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'file_type': 'pdf',
                'filename': os.path.basename(file_path)
            }
    
    def _parse_excel_document(self, file_path: str) -> Dict[str, Any]:
        """
        Extract data from Excel document
        
        Args:
            file_path: Path to Excel file
            
        Returns:
            Dictionary with extracted data and metadata
        """
        if not PANDAS_AVAILABLE:
            return {
                'success': False,
                'error': 'Excel support not available. Install pandas and openpyxl.',
                'file_type': 'excel',
                'filename': os.path.basename(file_path)
            }
        
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            
            all_text = []
            sheets_data = {}
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # Convert dataframe to text
                sheet_text = df.to_string(index=False)
                all_text.append(f"Sheet: {sheet_name}\n{sheet_text}")
                
                # Store structured data
                sheets_data[sheet_name] = df.to_dict('records')
            
            full_text = '\n\n'.join(all_text)
            
            # Detect sections
            sections = self._detect_sections(full_text)
            
            # Extract entities
            structured_data = None
            if self.entity_extractor:
                structured_data = self.entity_extractor.extract_structured_data(full_text)
            
            return {
                'success': True,
                'text': full_text,
                'sheets': sheets_data,
                'sheet_names': excel_file.sheet_names,
                'sections': sections,
                'structured_data': structured_data,
                'file_type': 'excel',
                'filename': os.path.basename(file_path)
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'file_type': 'excel',
                'filename': os.path.basename(file_path)
            }
    
    def _parse_image_document(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from image using OCR
        
        Args:
            file_path: Path to image file
            
        Returns:
            Dictionary with extracted text and metadata
        """
        if not self.ocr_handler:
            return {
                'success': False,
                'error': 'OCR not enabled',
                'file_type': 'image',
                'filename': os.path.basename(file_path)
            }
        
        try:
            ocr_result = self.ocr_handler.extract_text_from_image(file_path)
            
            if not ocr_result.get('success'):
                return {
                    'success': False,
                    'error': ocr_result.get('error', 'OCR failed'),
                    'file_type': 'image',
                    'filename': os.path.basename(file_path)
                }
            
            full_text = ocr_result.get('text', '')
            
            # Detect sections
            sections = self._detect_sections(full_text)
            
            # Extract entities
            structured_data = None
            if self.entity_extractor:
                structured_data = self.entity_extractor.extract_structured_data(full_text)
            
            return {
                'success': True,
                'text': full_text,
                'sections': sections,
                'structured_data': structured_data,
                'ocr_confidence': ocr_result.get('confidence', 0),
                'num_text_blocks': ocr_result.get('num_text_blocks', 0),
                'file_type': 'image',
                'filename': os.path.basename(file_path)
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'file_type': 'image',
                'filename': os.path.basename(file_path)
            }
    
    def _detect_sections(self, text: str) -> Dict[str, List[str]]:
        """
        Detect common sections in document
        
        Args:
            text: Document text
            
        Returns:
            Dictionary of section names and their content
        """
        sections = {
            'introduction': [],
            'requirements': [],
            'scope': [],
            'budget': [],
            'timeline': [],
            'other': []
        }
        
        # Common section headers
        section_patterns = {
            'introduction': r'(?i)(introduction|overview|background|about)',
            'requirements': r'(?i)(requirements?|specifications?|needs?)',
            'scope': r'(?i)(scope|deliverables?|objectives?)',
            'budget': r'(?i)(budget|cost|pricing|financial)',
            'timeline': r'(?i)(timeline|schedule|deadline|milestones?)'
        }
        
        lines = text.split('\n')
        current_section = 'other'
        
        for line in lines:
            line_stripped = line.strip()
            
            if not line_stripped:
                continue
            
            # Check if line is a section header
            section_found = False
            for section_name, pattern in section_patterns.items():
                if re.match(pattern, line_stripped):
                    current_section = section_name
                    section_found = True
                    break
            
            if not section_found:
                sections[current_section].append(line_stripped)
        
        return sections
    
    def extract_requirements(self, parsed_data: Dict[str, Any]) -> str:
        """
        Structure extracted text into requirements format
        
        Args:
            parsed_data: Dictionary from parse_document()
            
        Returns:
            Formatted requirements text
        """
        if not parsed_data.get('success'):
            raise ValueError(f"Document parsing failed: {parsed_data.get('error')}")
        
        text = parsed_data.get('text', '')
        
        # Clean up the text
        lines = text.split('\n')
        cleaned_lines = [line.strip() for line in lines if line.strip()]
        
        return '\n'.join(cleaned_lines)


def allowed_file(filename: str, allowed_extensions: set) -> bool:
    """
    Check if file has allowed extension
    
    Args:
        filename: Name of the file
        allowed_extensions: Set of allowed extensions
        
    Returns:
        True if file extension is allowed
    """
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in allowed_extensions
